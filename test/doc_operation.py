import os
import docx
import re


def debug():
    filename = "~/Desktop/abc.docx"
    doc = docx.Document(filename)
    # print(doc)

    tk1 = doc.paragraphs[3:44]  # 41
    pd1 = doc.paragraphs[45:227]  # 180
    dx1 = doc.paragraphs[228:493]  # 53

    tk2 = doc.paragraphs[510:538]  # 28
    pd2 = doc.paragraphs[539:574]  # 35
    dx2 = doc.paragraphs[575:615]  # 8
    ddx2 = doc.paragraphs[616:666]  # 10

    tk3 = doc.paragraphs[696:791]  # 95
    pd3 = doc.paragraphs[792:931]  # 139
    dx3 = doc.paragraphs[932:1117]  # 37
    ddx3 = doc.paragraphs[1119:1336]  # 43

    # tmp_doc = docx.Document()
    # for p in dx1 + dx2 + dx3:
    #     tmp_doc.add_paragraph(p.text)
    # tmp_doc.save("/Users/cxq/Desktop/danxuan.docx")

    # for p in pd3:
    #     print(f"{p.text}")

    # pattern = "^\d+\."
    # for p in ddx3:
    #     if re.match(pattern, p.text):
    #         print("-" * 72)
    #     print(f"{p.text}")

    # cnt = 0
    # for p in ddx3:
    #     cnt += 1
    #     print(f"=>{p.text}")
    #     if cnt % 5 == 0:
    #         print("-" * 72)

    # p3 = doc.paragraphs[3]
    # for each in p3.runs:
    #     print(f"[{each.text}, {each.underline}]", end='/')


def parse_fill_blank_paragraph(paragraph):
    content = list()
    answer = list()
    for each in paragraph.runs:
        if each.underline:
            answer.append(each.text)
            content.append("_" * len(each.text))
        else:
            content.append(each.text)

    return ''.join(content), ','.join(answer)


def get_fill_blank(doc):
    tk1 = doc.paragraphs[3:(3 + 41)]  # 41
    tk2 = doc.paragraphs[510:(510 + 28)]  # 28
    tk3 = doc.paragraphs[697:(697 + 95)]  # 95
    tk = tk1# + tk2 + tk3

    for paragraph in tk:
        print(f"ori: {paragraph.text}")
        content, answer = parse_fill_blank_paragraph(paragraph)
        print(f"que: {content}\n\tans: {answer}")


def parse_jude(paragraph):
    # answer = re.findall('^\d+\. \( (.*) \)', paragraph.text, re.S)
    text = paragraph.text
    s = text.find("(")
    e = text.find(")")
    answer = text[s + 1:e].strip()
    content = f"{text[:s]}(  ){text[e + 1:]}"
    return content, answer


def get_judge(doc):
    pd1 = doc.paragraphs[45:(45 + 180)]  # 180
    pd2 = doc.paragraphs[539:(539 + 35)]  # 35
    pd3 = doc.paragraphs[793:(793 + 139)]  # 139
    pd = pd1 + pd2 + pd3

    for paragraph in pd:
        print(f"ori: {paragraph.text}")
        content, answer = parse_jude(paragraph)
        print(f"gen: {content}\n\tans: {answer}")


def parse_single_select(paragraphs):
    topic = paragraphs[0].text
    s = topic.find("(")
    e = topic.find(")")
    answer = topic[s + 1:e].strip()

    topic_blank = re.sub("\([A-Z]\)", "( )", topic)
    candidates = '\n'.join([each.text for each in paragraphs[1:]])
    content = f"{topic_blank}\n{candidates}"

    return content, answer


def get_single_select(doc):
    dx1 = doc.paragraphs[226:(226 + 53 * 5)]  # 53
    dx2 = doc.paragraphs[575:(575 + 8 * 5)]  # 8
    dx3 = doc.paragraphs[933:(933 + 37 * 5)]  # 37

    dx = dx1 + dx2 + dx3

    # for p in dx:
    #     print(p.text)

    pattern = "^\d+\."
    topic_idx = []
    for k, paragraph in enumerate(dx):
        if re.match(pattern, paragraph.text):
            topic_idx.append(k)
    topic_idx.append(len(dx))

    for i in range(len(topic_idx) - 1):
        print(f"ori: {dx[topic_idx[i]].text}")
        content, answer = parse_single_select(dx[topic_idx[i]:topic_idx[i + 1]])
        print(f"gen: {content}\n\tans: {answer}")


def parse_multiple_select(paragraphs):
    topic = paragraphs[0].text
    s = topic.find("(")
    e = topic.find(")")
    answer = topic[s + 1:e].strip()

    topic_blank = re.sub("\([A-Z]+\)", "( )", topic)
    candidates = '\n'.join([each.text for each in paragraphs[1:]])
    content = f"{topic_blank}\n{candidates}"

    return content, answer


def get_multiple_select(doc):
    dx2 = doc.paragraphs[616:(616 + 10 * 5)]  # 10
    dx3 = doc.paragraphs[1120:(1120 + 43 * 5 + 2)]  # 43
    dx = dx2 + dx3

    # for p in dx:
    #     print(p.text)

    pattern = "^\d+\."
    topic_idx = []
    for k, paragraph in enumerate(dx):
        if re.match(pattern, paragraph.text):
            topic_idx.append(k)
    topic_idx.append(len(dx))

    for i in range(len(topic_idx) - 1):
        print(f"ori: {dx[topic_idx[i]].text}")
        content, answer = parse_multiple_select(dx[topic_idx[i]:topic_idx[i + 1]])
        print(f"gen: {content}\n\tans: {answer}")


def parse_short_answer(paragraphs):
    content = paragraphs[0].text
    answer = "\n".join([each.text for each in paragraphs[1:]])
    return content, answer


def get_short_answer(doc):
    jd1 = doc.paragraphs[499:(499 + 4 * 2)]
    jd2 = doc.paragraphs[669:694]
    jd3 = doc.paragraphs[1338:1374]
    jd = jd1 + jd2 + jd3

    for paragraph in jd:
        print(f"{paragraph.text}")

    pattern = "^\d+\."
    topic_idx = []
    for k, paragraph in enumerate(jd):
        if re.match(pattern, paragraph.text):
            topic_idx.append(k)
    topic_idx.append(len(jd))

    for i in range(len(topic_idx) - 1):
        print(f"ori: {jd[topic_idx[i]].text}")
        content, answer = parse_short_answer(jd[topic_idx[i]:topic_idx[i + 1]])
        print(f"gen: {content}\n\tans: {answer}")


def main():
    filename = "/Users/cxq/Desktop/abc.docx"
    doc = docx.Document(filename)
    get_fill_blank(doc)
    # get_judge(doc)
    # get_single_select(doc)
    # get_multiple_select(doc)
    # get_short_answer(doc)


if __name__ == '__main__':
    main()
