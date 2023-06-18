import docx
import re
import random


class BMTest(object):
    def __init__(self, file_name, tk=10, pd=10, dx=10, ddx=10, jd=3):
        self.doc = docx.Document(file_name)
        self.tk = tk
        self.pd = pd
        self.dx = dx
        self.ddx = ddx
        self.jd = jd

    @staticmethod
    def parse_fill_blank_paragraph(paragraph):
        content = list()
        answer = list()
        for each in paragraph.runs:
            if each.underline:
                answer.append(each.text)
                content.append("_" * len(each.text))
            else:
                content.append(each.text)

        return ''.join(content), ','.join(answer)

    def get_fill_blank(self):
        tk1 = self.doc.paragraphs[3:(3 + 41)]  # 41
        tk2 = self.doc.paragraphs[510:(510 + 28)]  # 28
        tk3 = self.doc.paragraphs[697:(697 + 95)]  # 95
        tk = tk1 + tk2 + tk3

        return [self.parse_fill_blank_paragraph(paragraph) for paragraph in tk]

    @staticmethod
    def parse_jude(paragraph):
        text = paragraph.text
        s = text.find("(")
        e = text.find(")")
        answer = text[s + 1:e].strip()
        content = f"{text[:s]}(  ){text[e + 1:]}"
        return content, answer

    def get_judge(self):
        pd1 = self.doc.paragraphs[45:(45 + 180)]  # 180
        pd2 = self.doc.paragraphs[539:(539 + 35)]  # 35
        pd3 = self.doc.paragraphs[793:(793 + 139)]  # 139
        pd = pd1 + pd2 + pd3

        return [self.parse_jude(paragraph) for paragraph in pd]

    @staticmethod
    def parse_single_select(paragraphs):
        topic = paragraphs[0].text
        s = topic.find("(")
        e = topic.find(")")
        answer = topic[s + 1:e].strip()

        topic_blank = re.sub("\([A-Z]\)", "( )", topic)
        candidates = '\n'.join([each.text for each in paragraphs[1:]])
        content = f"{topic_blank}\n{candidates}"

        return content, answer

    def get_single_select(self):
        dx1 = self.doc.paragraphs[226:(226 + 53 * 5)]  # 53
        dx2 = self.doc.paragraphs[575:(575 + 8 * 5)]  # 8
        dx3 = self.doc.paragraphs[933:(933 + 37 * 5)]  # 37
        dx = dx1 + dx2 + dx3

        topic_idx = self.get_topic_idx(dx)

        return [self.parse_single_select(dx[topic_idx[i]:topic_idx[i + 1]]) for i in range(len(topic_idx) - 1)]

    @staticmethod
    def parse_multiple_select(paragraphs):
        topic = paragraphs[0].text
        s = topic.find("(")
        e = topic.find(")")
        answer = topic[s + 1:e].strip()

        topic_blank = re.sub("\([A-Z]+\)", "( )", topic)
        candidates = '\n'.join([each.text for each in paragraphs[1:]])
        content = f"{topic_blank}\n{candidates}"

        return content, answer

    def get_multiple_select(self):
        dx2 = self.doc.paragraphs[616:(616 + 10 * 5)]  # 10
        dx3 = self.doc.paragraphs[1120:(1120 + 43 * 5 + 2)]  # 43
        dx = dx2 + dx3

        topic_idx = self.get_topic_idx(dx)

        return [self.parse_multiple_select(dx[topic_idx[i]:topic_idx[i + 1]]) for i in range(len(topic_idx) - 1)]

    @staticmethod
    def parse_short_answer(paragraphs):
        content = paragraphs[0].text
        answer = "\n".join([each.text for each in paragraphs[1:]])
        return content, answer

    def get_short_answer(self):
        jd1 = self.doc.paragraphs[499:(499 + 4 * 2)]
        jd2 = self.doc.paragraphs[669:694]
        jd3 = self.doc.paragraphs[1338:1374]
        jd = jd1 + jd2 + jd3

        topic_idx = self.get_topic_idx(jd)

        return [self.parse_short_answer(jd[topic_idx[i]:topic_idx[i + 1]]) for i in range(len(topic_idx) - 1)]

    @staticmethod
    def get_topic_idx(paragraphs):
        pattern = "^\d+\."
        topic_idx = list()
        for k, paragraph in enumerate(paragraphs):
            if re.match(pattern, paragraph.text):
                topic_idx.append(k)
        topic_idx.append(len(paragraphs))

        return topic_idx

    @staticmethod
    def sample_exam(lst, k):

        return [(re.sub("^\d+", str(k + 1), content), answer) for k, (content, answer) in
                enumerate(random.sample(lst, k))]

    def gen_examination(self):
        blanks = self.get_fill_blank()
        judges = self.get_judge()
        single_select = self.get_single_select()
        multiple_select = self.get_multiple_select()
        short_answer = self.get_short_answer()

        exams = self.sample_exam(blanks, self.tk) + self.sample_exam(judges, self.pd) \
                + self.sample_exam(single_select, self.dx) + self.sample_exam(multiple_select, self.ddx) \
                + self.sample_exam(short_answer, self.jd)

        return exams

    def run(self):
        exams = self.gen_examination()
        for content, answer in exams:
            print(f"{content}\nans: {answer}\n---------")


def main():
    filename = "/Users/cxq/Desktop/abc.docx"
    ex_test = BMTest(file_name=filename)
    ex_test.run()


if __name__ == '__main__':
    main()
